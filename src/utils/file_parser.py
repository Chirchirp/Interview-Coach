"""File parser - PDF, DOCX, TXT extraction."""
from __future__ import annotations
import io, re


def extract_text(uploaded_file) -> tuple[str, str]:
    """Returns (text, error). error is empty on success."""
    if uploaded_file is None:
        return "", "No file provided."
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    try:
        if name.endswith(".txt"):
            for enc in ["utf-8","latin-1","cp1252"]:
                try: return data.decode(enc), ""
                except: pass
            return data.decode("utf-8", errors="replace"), ""
        elif name.endswith(".pdf"):
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(data))
                text = "\n".join(p.extract_text() or "" for p in reader.pages).strip()
                return (text, "") if text else ("", "PDF has no extractable text.")
            except ImportError:
                return "", "pypdf not installed. Run: pip install pypdf"
        elif name.endswith(".docx"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(data))
                paras = [p.text for p in doc.paragraphs if p.text.strip()]
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            if cell.text.strip(): paras.append(cell.text.strip())
                text = "\n".join(paras).strip()
                return (text, "") if text else ("", "DOCX has no extractable text.")
            except ImportError:
                return "", "python-docx not installed. Run: pip install python-docx"
        else:
            return "", f"Unsupported type: {name.split('.')[-1].upper()}"
    except Exception as e:
        return "", f"Parse error: {e}"


def clean(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()
